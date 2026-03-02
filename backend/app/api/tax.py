from fastapi import APIRouter, Depends, HTTPException, Response
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from app.db.session import get_session
from app.api.deps import get_current_user
from app.db.models import User, TaxLedger, Relationship
import io
import csv
from datetime import datetime
from weasyprint import HTML
import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

router = APIRouter(prefix="/api/tax", tags=["tax"])

def generate_tax_ledger_html(user: User, ledger: TaxLedger, relationships: dict) -> str:
    """Generate a clean HTML representation of the tax ledger."""
    current_year = f"{datetime.utcnow().year}"
    
    q_data = ledger.q1_q4_json or {"q1": 0, "q2": 0, "q3": 0, "q4": 0}
    rel_data = ledger.per_relationship_json or {}
    
    rel_rows = ""
    for rel_id, amount in rel_data.items():
        label = relationships.get(rel_id, rel_id)
        rel_rows += f"""
        <tr>
            <td style="padding: 10px; border-bottom: 1px solid #eee; font-size: 14px;">{label}</td>
            <td style="padding: 10px; border-bottom: 1px solid #eee; text-align: right; font-size: 14px;">${amount:,.2f}</td>
        </tr>
        """
        
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <style>
            body {{ font-family: 'Helvetica Neue', Arial, sans-serif; color: #1A1A18; line-height: 1.6; padding: 50px; }}
            .header {{ border-bottom: 2px solid #D4A017; padding-bottom: 20px; margin-bottom: 40px; }}
            .title {{ color: #D4A017; font-size: 32px; font-weight: bold; margin: 0; }}
            .meta {{ font-size: 12px; color: #8A8880; margin-top: 5px; }}
            .summary {{ display: flex; justify-content: space-between; margin-bottom: 40px; background: #FAF9F6; padding: 25px; border-radius: 12px; }}
            .stat {{ flex: 1; text-align: center; }}
            .stat-label {{ font-size: 11px; text-transform: uppercase; color: #8A8880; font-weight: bold; margin-bottom: 5px; }}
            .stat-value {{ font-size: 24px; font-weight: bold; color: #1A1A18; }}
            .quarterly {{ margin-bottom: 40px; }}
            .section-title {{ font-size: 18px; font-weight: bold; margin-bottom: 15px; border-bottom: 1px solid #E8E6E0; padding-bottom: 8px; }}
            table {{ width: 100%; border-collapse: collapse; }}
            th {{ text-align: left; padding: 10px; background: #F5F3EE; font-size: 12px; text-transform: uppercase; color: #8A8880; }}
            .footer {{ margin-top: 50px; font-size: 10px; color: #8A8880; text-align: center; border-top: 1px solid #E8E6E0; padding-top: 20px; }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1 class="title">Tax Ledger Report</h1>
            <div class="meta">Generated for {user.email} &bull; Tax Year {current_year} &bull; {datetime.utcnow().strftime('%Y-%m-%d')}</div>
        </div>

        <div class="summary">
            <div class="stat">
                <div class="stat-label">Total Earned</div>
                <div class="stat-value">${ledger.total_earned:,.2f}</div>
            </div>
            <div class="stat">
                <div class="stat-label">Tax Collected</div>
                <div class="stat-value">${ledger.total_tax_collected:,.2f}</div>
            </div>
            <div class="stat">
                <div class="stat-label">Estimated Owed (25%)</div>
                <div class="stat-value" style="color: #C0392B;">${(ledger.total_earned * 0.25):,.2f}</div>
            </div>
        </div>

        <div class="quarterly">
            <div class="section-title">Quarterly Performance</div>
            <table>
                <thead>
                    <tr><th>Quarter</th><th style="text-align: right;">Amount</th></tr>
                </thead>
                <tbody>
                    <tr><td>Q1 (Jan-Mar)</td><td style="text-align: right;">${q_data.get('q1', 0):,.2f}</td></tr>
                    <tr><td>Q2 (Apr-Jun)</td><td style="text-align: right;">${q_data.get('q2', 0):,.2f}</td></tr>
                    <tr><td>Q3 (Jul-Sep)</td><td style="text-align: right;">${q_data.get('q3', 0):,.2f}</td></tr>
                    <tr><td>Q4 (Oct-Dec)</td><td style="text-align: right;">${q_data.get('q4', 0):,.2f}</td></tr>
                </tbody>
            </table>
        </div>

        <div>
            <div class="section-title">Income Sources</div>
            <table>
                <thead>
                    <tr><th>Relationship / Client</th><th style="text-align: right;">Amount</th></tr>
                </thead>
                <tbody>
                    {rel_rows if rel_rows else '<tr><td colspan="2" style="text-align: center; color: #8A8880; padding: 20px;">No transaction data found.</td></tr>'}
                </tbody>
            </table>
        </div>

        <div class="footer">
            Generated by INVOQ — Relationship-First Infrastructure for Modern Contractors.
        </div>
    </body>
    </html>
    """
    return html

@router.get("/ledger", response_model=dict)
async def get_tax_ledger(
    user: User = Depends(get_current_user),
    session: AsyncSession = Depends(get_session)
):
    """Get the current user's tax ledger for the current tax year."""
    if user.subscription_status != "pro":
        raise HTTPException(
            status_code=403, 
            detail="Tax Ledger is a Pro feature. Please upgrade to access tax tracking and exports."
        )

    current_year = f"{datetime.utcnow().year}"
    result = await session.execute(
        select(TaxLedger).where(
            TaxLedger.user_id == user.id,
            TaxLedger.tax_year == current_year
        )
    )
    ledger = result.scalars().first()
    
    if not ledger:
        # Return empty shell if no data yet
        return {
            "total_earned": 0,
            "total_tax_collected": 0,
            "q1_q4_json": {"q1": 0, "q2": 0, "q3": 0, "q4": 0},
            "per_relationship_json": {}
        }
    
    return ledger

@router.get("/export/csv")
async def export_tax_csv(
    user: User = Depends(get_current_user),
    session: AsyncSession = Depends(get_session)
):
    """Export the tax ledger as a CSV file."""
    if user.subscription_status != "pro":
        raise HTTPException(
            status_code=403, 
            detail="Tax Ledger is a Pro feature. Please upgrade to access tax tracking and exports."
        )

    current_year = f"{datetime.utcnow().year}"
    result = await session.execute(
        select(TaxLedger).where(
            TaxLedger.user_id == user.id,
            TaxLedger.tax_year == current_year
        )
    )
    ledger = result.scalars().first()
    
    output = io.StringIO()
    writer = csv.writer(output)
    
    writer.writerow(["Tax Ledger Export", f"Year: {current_year}"])
    writer.writerow(["User", user.email])
    writer.writerow([])
    
    if not ledger:
        writer.writerow(["No tax data available for this period."])
    else:
        writer.writerow(["Summary"])
        writer.writerow(["Total Earned", ledger.total_earned])
        writer.writerow(["Total Tax Collected", ledger.total_tax_collected])
        writer.writerow([])
        
        writer.writerow(["Quarterly Breakdown"])
        q_data = ledger.q1_q4_json or {}
        writer.writerow(["Quarter", "Amount"])
        for q in ["q1", "q2", "q3", "q4"]:
            writer.writerow([q.upper(), q_data.get(q, 0)])
        writer.writerow([])
        
        writer.writerow(["Client / Relationship Breakdown"])
        rel_data = ledger.per_relationship_json or {}
        writer.writerow(["Relationship ID", "Amount"]) # We could fetch names, but keeping it simple for V1
        for rel_id, amount in rel_data.items():
            # Try to fetch relationship name for better UX
            rel_result = await session.execute(select(Relationship).where(Relationship.id == rel_id))
            rel = rel_result.scalars().first()
            label = rel.recipient_email if rel else rel_id
            writer.writerow([label, amount])
            
    return Response(
        content=output.getvalue(), # Changed 'content' to 'output.getvalue()'
        media_type="text/csv",
        headers={
            "Content-Disposition": f'attachment; filename="tax_ledger_{current_year}.csv"'
        }
    )

@router.get("/export/pdf")
async def export_tax_pdf(
    user: User = Depends(get_current_user),
    session: AsyncSession = Depends(get_session)
):
    """Export the tax ledger as a nicely formatted PDF."""
    if user.subscription_status != "pro":
        raise HTTPException(
            status_code=403, 
            detail="Tax Ledger is a Pro feature. Please upgrade to access tax tracking and exports."
        )

    current_year = f"{datetime.utcnow().year}"
    result = await session.execute(
        select(TaxLedger).where(
            TaxLedger.user_id == user.id,
            TaxLedger.tax_year == current_year
        )
    )
    ledger = result.scalars().first()
    
    if not ledger:
        ledger = TaxLedger(
            user_id=user.id,
            tax_year=current_year,
            total_earned=0,
            total_tax_collected=0,
            q1_q4_json={"q1": 0, "q2": 0, "q3": 0, "q4": 0},
            per_relationship_json={}
        )

    # Fetch relationship names for the report
    rel_names = {}
    rel_data = ledger.per_relationship_json or {}
    for rel_id in rel_data.keys():
        rel_result = await session.execute(select(Relationship).where(Relationship.id == rel_id))
        rel = rel_result.scalars().first()
        if rel:
            rel_names[rel_id] = rel.recipient_email

    html_content = generate_tax_ledger_html(user, ledger, rel_names)
    pdf_bytes = HTML(string=html_content).write_pdf()
    
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="tax_ledger_{current_year}.pdf"'
        }
    )

@router.get("/export/excel")
async def export_tax_excel(
    user: User = Depends(get_current_user),
    session: AsyncSession = Depends(get_session)
):
    """Export the tax ledger as an Excel file."""
    if user.subscription_status != "pro":
        raise HTTPException(status_code=403, detail="Tax Ledger is a Pro feature.")

    current_year = f"{datetime.utcnow().year}"
    result = await session.execute(
        select(TaxLedger).where(
            TaxLedger.user_id == user.id,
            TaxLedger.tax_year == current_year
        )
    )
    ledger = result.scalars().first()
    if not ledger:
        return Response(status_code=404, content="No data")

    # Fetch relationships
    rel_data = ledger.per_relationship_json or {}
    rel_names = {}
    for rel_id in rel_data.keys():
        r = await session.execute(select(Relationship).where(Relationship.id == rel_id))
        if rel := r.scalars().first():
            rel_names[rel_id] = rel.recipient_email

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"Tax Ledger {current_year}"

    ws.append(["Tax Ledger Report", current_year])
    ws.append([])
    ws.append(["Total Earned", ledger.total_earned])
    ws.append(["Tax Collected", ledger.total_tax_collected])
    ws.append(["Estimated Owed (25%)", ledger.total_earned * 0.25])
    ws.append([])
    
    ws.append(["Quarter", "Amount"])
    q_data = ledger.q1_q4_json or {}
    for q in ["q1", "q2", "q3", "q4"]:
        ws.append([q.upper(), q_data.get(q, 0)])
        
    ws.append([])
    ws.append(["Client / Relationship", "Amount"])
    for rel_id, amount in rel_data.items():
        ws.append([rel_names.get(rel_id, rel_id), amount])

    output = io.BytesIO()
    wb.save(output)
    
    return Response(
        content=output.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="tax_ledger_{current_year}.xlsx"'}
    )

@router.get("/export/docx")
async def export_tax_docx(
    user: User = Depends(get_current_user),
    session: AsyncSession = Depends(get_session)
):
    """Export the tax ledger as a DOCX file."""
    if user.subscription_status != "pro":
        raise HTTPException(status_code=403, detail="Tax Ledger is a Pro feature.")

    current_year = f"{datetime.utcnow().year}"
    result = await session.execute(
        select(TaxLedger).where(
            TaxLedger.user_id == user.id,
            TaxLedger.tax_year == current_year
        )
    )
    ledger = result.scalars().first()
    if not ledger:
        return Response(status_code=404, content="No data")

    rel_data = ledger.per_relationship_json or {}
    rel_names = {}
    for rel_id in rel_data.keys():
        r = await session.execute(select(Relationship).where(Relationship.id == rel_id))
        if rel := r.scalars().first():
            rel_names[rel_id] = rel.recipient_email

    doc = Document()
    head = doc.add_heading(f"Tax Ledger Report - {current_year}", 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated for: {user.email}\nDate: {datetime.utcnow().strftime('%Y-%m-%d')}")

    doc.add_heading("Summary", level=1)
    p = doc.add_paragraph()
    p.add_run("Total Earned: ").bold = True
    p.add_run(f"${ledger.total_earned:,.2f}\n")
    p.add_run("Total Tax Collected: ").bold = True
    p.add_run(f"${ledger.total_tax_collected:,.2f}\n")
    p.add_run("Estimated Owed (25%): ").bold = True
    r = p.add_run(f"${ledger.total_earned * 0.25:,.2f}")
    r.font.color.rgb = RGBColor(192, 57, 43)

    doc.add_heading("Quarterly Performance", level=1)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Quarter"
    hdr_cells[1].text = "Amount"
    q_data = ledger.q1_q4_json or {}
    for q in ["q1", "q2", "q3", "q4"]:
        row_cells = table.add_row().cells
        row_cells[0].text = q.upper()
        row_cells[1].text = f"${q_data.get(q, 0):,.2f}"

    doc.add_heading("Income Sources", level=1)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Client"
    hdr_cells[1].text = "Amount"
    for rel_id, amount in rel_data.items():
        row_cells = table.add_row().cells
        row_cells[0].text = rel_names.get(rel_id, rel_id)
        row_cells[1].text = f"${amount:,.2f}"

    output = io.BytesIO()
    doc.save(output)
    
    return Response(
        content=output.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="tax_ledger_{current_year}.docx"'}
    )
