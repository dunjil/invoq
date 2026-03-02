import io
import re
from typing import Optional, List
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class DOCXService:
    @staticmethod
    def generate_document(data: dict, title: str = "Agreement") -> io.BytesIO:
        """
        Generate a professional DOCX document from document data.
        Expected keys in data: title, document_number, effective_date, from_name, to_name, body_text, primary_color.
        """
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        style.font.name = 'Helvetica'
        style.font.size = Pt(11)
        
        # Accent Color
        primary_color_hex = data.get('primary_color', '#D4A017').lstrip('#')
        accent_color = RGBColor.from_string(primary_color_hex)

        # Header: Title
        header_title = doc.add_paragraph()
        header_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = header_title.add_run(data.get('title', title).upper())
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = accent_color
        
        # Header: Metadata
        meta = doc.add_paragraph()
        meta.add_run(f"Ref: {data.get('document_number', 'N/A')}\n")
        meta.add_run(f"Effective Date: {data.get('effective_date', 'Upon Signature')}")
        meta.style.font.size = Pt(9)
        meta.style.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        
        # Divider
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(20)
        
        # Parties Section
        table = doc.add_table(rows=1, cols=2)
        table.width = Inches(6)
        
        # Party 1
        cell1 = table.rows[0].cells[0]
        p1 = cell1.paragraphs[0]
        run1 = p1.add_run("PARTY 1 (PROVIDER)\n")
        run1.bold = True
        run1.font.size = Pt(8)
        run1.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        p1.add_run(data.get('from_name', ''))
        
        # Party 2
        cell2 = table.rows[0].cells[1]
        p2 = cell2.paragraphs[0]
        run2 = p2.add_run("PARTY 2 (CLIENT)\n")
        run2.bold = True
        run2.font.size = Pt(8)
        run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        p2.add_run(data.get('to_name', ''))
        
        doc.add_paragraph().paragraph_format.space_after = Pt(20)

        # Body Text (Markdown Parsing)
        body_text = data.get('body_text', '')
        DOCXService._parse_markdown(doc, body_text)

        # Signatures
        doc.add_page_break()
        doc.add_heading('Signatures', level=2)
        sig_table = doc.add_table(rows=1, cols=2)
        
        # Issuer Sig
        if data.get('include_issuer_signature', True):
            c1 = sig_table.rows[0].cells[0]
            p = c1.paragraphs[0]
            p.add_run("\n\n__________________________\n")
            p.add_run(f"Authorized Signature ({data.get('from_name', '')})")
            
        # Recipient Sig
        if data.get('include_recipient_signature', True):
            c2 = sig_table.rows[0].cells[1]
            p = c2.paragraphs[0]
            p.add_run("\n\n__________________________\n")
            p.add_run(f"Client Signature ({data.get('to_name', '')})")

        # Footer
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Generated via Invoq.app — Secure Document Management")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        target = io.BytesIO()
        doc.save(target)
        target.seek(0)
        return target

    @staticmethod
    def _parse_markdown(doc: Document, text: str):
        """Simple markdown parser for python-docx."""
        lines = text.split('\n')
        in_list = False
        in_table = False
        table_data = []

        for line in lines:
            line = line.strip()
            if not line:
                if in_table:
                    DOCXService._add_table_to_doc(doc, table_data)
                    table_data = []
                    in_table = False
                continue

            # Headers
            if line.startswith('### '):
                h = doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                h = doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                h = doc.add_heading(line[2:], level=1)
            
            # Lists
            elif line.startswith('* ') or line.startswith('- '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                p = doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
            
            # Tables
            elif line.startswith('|') and '|' in line[1:]:
                in_table = True
                # Clean up the line and split by |
                cells = [c.strip() for c in line.split('|') if c.strip()]
                if not all(c == '-' or c.startswith('---') for c in cells): # Skip separator lines
                    table_data.append(cells)
            
            else:
                if in_table:
                    DOCXService._add_table_to_doc(doc, table_data)
                    table_data = []
                    in_table = False
                
                # Normal paragraph with basic bold parsing
                p = doc.add_paragraph()
                DOCXService._add_formatted_text(p, line)

        if in_table:
            DOCXService._add_table_to_doc(doc, table_data)

    @staticmethod
    def _add_formatted_text(paragraph, text):
        """Handle basic bold (**text**) and italic (*text*) parsing."""
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                paragraph.add_run(part)

    @staticmethod
    def _add_table_to_doc(doc, data):
        if not data: return
        rows = len(data)
        cols = max(len(row) for row in data)
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        for i, row_data in enumerate(data):
            for j, val in enumerate(row_data):
                if j < len(table.rows[i].cells):
                    table.rows[i].cells[j].text = val
